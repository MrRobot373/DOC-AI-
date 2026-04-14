import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from doc_parser import parse_document
from review_engine import (
    _check_heading_hierarchy,
    _check_section_number_continuity,
    _check_toc_heading_sync,
)


def _ensure_paragraph_style(doc, style_name):
    """Create the paragraph style when the default template does not provide it."""
    if style_name in [style.name for style in doc.styles]:
        return

    style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]


class TocValidationTests(unittest.TestCase):
    def _parse_fixture(self, builder):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "fixture.docx"
            doc = Document()
            for style_name in ("TOC 1", "TOC 2", "TOC 3", "TOC Heading"):
                _ensure_paragraph_style(doc, style_name)
            builder(doc)
            doc.save(path)
            return parse_document(str(path))

    def test_parser_extracts_toc_entries_without_promoting_them_to_headings(self):
        def builder(doc):
            doc.add_paragraph("Table of Contents", style="TOC Heading")
            doc.add_paragraph("1 Introduction\t1", style="TOC 1")
            doc.add_paragraph("1.1 Scope\t2", style="TOC 2")
            doc.add_paragraph("2 Design\t3", style="TOC 1")

            doc.add_heading("1 Introduction", 1)
            doc.add_heading("1.1 Scope", 2)
            doc.add_heading("2 Design", 1)

        parsed = self._parse_fixture(builder)

        self.assertEqual(parsed["toc"]["title"], "Table of Contents")
        self.assertEqual(
            [entry["text"] for entry in parsed["toc"]["entries"]],
            ["1 Introduction", "1.1 Scope", "2 Design"],
        )
        self.assertEqual(
            [heading["text"] for heading in parsed["headings"]],
            ["1 Introduction", "1.1 Scope", "2 Design"],
        )

    def test_toc_sync_flags_mismatched_title_and_missing_heading(self):
        def builder(doc):
            doc.add_paragraph("Table of Contents", style="TOC Heading")
            doc.add_paragraph("1 Introduction\t1", style="TOC 1")
            doc.add_paragraph("2 Wrong Title\t2", style="TOC 1")

            doc.add_heading("1 Introduction", 1)
            doc.add_heading("2 Design", 1)
            doc.add_heading("3 Conclusion", 1)

        parsed = self._parse_fixture(builder)
        findings = _check_toc_heading_sync(parsed)
        comments = [finding["comment"] for finding in findings]

        self.assertTrue(
            any("does not match the actual heading text for section 2" in comment for comment in comments)
        )
        self.assertTrue(
            any("Heading '3 Conclusion' appears in the document body but is missing from the TOC." == comment for comment in comments)
        )

    def test_section_number_continuity_flags_gap_and_duplicate(self):
        def builder(doc):
            doc.add_heading("1 Scope", 1)
            doc.add_heading("1.1 Inputs", 2)
            doc.add_heading("1.3 Outputs", 2)
            doc.add_heading("1.3 Timing", 2)

        parsed = self._parse_fixture(builder)
        findings = _check_section_number_continuity(parsed)
        comments = [finding["comment"] for finding in findings]

        self.assertTrue(
            any("Missing section number(s): 1.2." in comment for comment in comments)
        )
        self.assertTrue(
            any("Duplicate section number '1.3'" in comment for comment in comments)
        )

    def test_heading_hierarchy_flags_orphan_subsection(self):
        def builder(doc):
            doc.add_heading("1 Scope", 1)
            doc.add_heading("1.1.1 Deep Details", 3)

        parsed = self._parse_fixture(builder)
        findings = _check_heading_hierarchy(parsed)
        comments = [finding["comment"] for finding in findings]

        self.assertTrue(
            any("Orphan heading '1.1.1 Deep Details'" in comment for comment in comments)
        )

    def test_toc_sync_zero_findings_on_clean_doc(self):
        def builder(doc):
            doc.add_paragraph("Table of Contents", style="TOC Heading")
            doc.add_paragraph("1 Introduction\t1", style="TOC 1")
            doc.add_paragraph("1.1 Scope\t2", style="TOC 2")
            doc.add_paragraph("2 Design\t3", style="TOC 1")

            doc.add_heading("1 Introduction", 1)
            doc.add_heading("1.1 Scope", 2)
            doc.add_heading("2 Design", 1)

        parsed = self._parse_fixture(builder)
        findings = _check_toc_heading_sync(parsed)

        self.assertEqual(len(findings), 0, "Expected 0 findings on perfectly matched TOC and headings")


if __name__ == "__main__":
    unittest.main()
