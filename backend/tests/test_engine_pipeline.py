"""
End-to-end review pipeline tests using a deterministic mock Ollama client.

These verify the Phase 0 wiring without a live LLM:
  - the full pipeline runs (local checks -> LLM passes -> grounding -> critic -> dedupe)
  - output is deterministic (same input -> identical findings)
  - hallucinated findings (evidence not in the document) are dropped
  - per-pass status is reported
"""

import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from doc_parser import parse_document
from review_engine import review_document


class FakeOllamaClient:
    """Deterministic stand-in for the Ollama client."""

    def __init__(self):
        self.calls = 0

    def chat(self, model, messages, format=None, options=None):
        self.calls += 1
        prompt = messages[0]["content"]
        # Critic pass: echo "keep" for every candidate id it was given.
        if "STRICT QA reviewer" in prompt:
            import re
            ids = sorted(set(int(n) for n in re.findall(r'"id":\s*(\d+)', prompt)))
            verdicts = [{"id": i, "keep": True, "confidence": 0.8, "reason": "ok"} for i in ids]
            return {"message": {"content": json.dumps(verdicts)}}
        # Findings pass: one grounded finding (quote exists) + one hallucination.
        findings = [
            {
                "category": "GRAMMAR_SPELLING", "severity": "MINOR",
                "page": "-", "section": "Introduction",
                "comment": "Awkward phrasing near the scope statement.",
                "fix": "Rephrase.", "fix_type": "MANUAL",
                "evidence": "Scope of this document",
            },
            {
                "category": "LOGICAL_CONSISTENCY", "severity": "CRITICAL",
                "page": "-", "section": "ALL",
                "comment": "Contradiction about the flux capacitor subsystem.",
                "fix": "Investigate.", "fix_type": "MANUAL",
                "evidence": "the flux capacitor overdrive subsystem is misconfigured",
            },
        ]
        return {"message": {"content": json.dumps(findings)}}


def _build_doc(path):
    doc = Document()
    doc.add_heading("1 Introduction", 1)
    doc.add_paragraph("Scope of this document is the ACC power supply design.")
    doc.add_heading("2 Design", 1)
    doc.add_paragraph("The output voltage is 3.3 V nominal under all conditions.")
    doc.save(path)


class EnginePipelineTests(unittest.TestCase):
    def _run(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "fixture.docx"
            _build_doc(path)
            parsed = parse_document(str(path))
            status = {}
            findings = review_document(
                FakeOllamaClient(), "fake-model", parsed,
                review_mode="pro", status_out=status,
            )
            return findings, status

    def test_pipeline_runs_and_grounds(self):
        # Recall-first: in Pro mode NOTHING is dropped. The ungrounded "flux
        # capacitor" finding is KEPT but flagged grounded=False at low confidence,
        # while the grounded finding stays high-confidence.
        findings, status = self._run()
        flux = [f for f in findings if "flux capacitor" in f["comment"].lower()]
        self.assertTrue(flux, "Ungrounded findings should be kept (recall-first), not dropped")
        self.assertFalse(flux[0].get("grounded", True), "Ungrounded finding should be flagged grounded=False")
        self.assertLessEqual(flux[0]["confidence"], 0.4, "Ungrounded finding should be low-confidence")

        grounded = [f for f in findings if "scope statement" in f["comment"].lower()]
        self.assertTrue(grounded, "Expected the grounded finding to survive")
        # Every finding carries the new schema fields.
        for f in findings:
            self.assertIn("evidence", f)
            self.assertIn("para_index", f)
            self.assertIn("confidence", f)
        self.assertIn("passes", status)

    def test_max_mode_drops_low_confidence(self):
        # Max mode IS the precision cut: the ungrounded low-confidence finding is
        # removed, the grounded one stays.
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "fixture.docx"
            _build_doc(path)
            parsed = parse_document(str(path))
            findings = review_document(FakeOllamaClient(), "fake-model", parsed, review_mode="max")
        comments = " ".join(f["comment"] for f in findings).lower()
        self.assertNotIn("flux capacitor", comments)

    def test_pipeline_is_deterministic(self):
        a, _ = self._run()
        b, _ = self._run()
        norm = lambda fs: [(f["category"], f["severity"], f["comment"]) for f in fs]
        self.assertEqual(norm(a), norm(b), "Same input must produce identical findings")


if __name__ == "__main__":
    unittest.main()
