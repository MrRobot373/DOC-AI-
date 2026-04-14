"""
Document Parser Module
Extracts structured content from .docx files including text, headings,
tables, images, and formatting metadata.
"""

import os
import io
import re
import base64
import zipfile
import shutil
from docx import Document
from docx.text.paragraph import Paragraph
import openpyxl
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import openpyxl


def parse_document(filepath):
    """
    Parse a .docx file and extract all content in a structured format.
    """
    try:
        doc = Document(filepath)
    except KeyError as e:
        if "NULL" in str(e):
            print(f"Sanitizing corrupt DOCX relationships for {filepath}")
            _sanitize_docx(filepath)
            doc = Document(filepath) # Re-attempt parsing after sanitization
        else:
            raise e
    except Exception as e:
        # Fallback to a very simple parse if something fails
        return {
            "statistics": {"total_words": 0, "total_sections": 0, "total_tables": 0, "total_images": 0},
            "sections": [{"title": "Error", "level": 1, "paragraphs": [{"text": f"Error parsing document: {str(e)}", "heading_level": 0, "alignment": "LEFT", "runs": [], "has_image": False}]}]
        }
    result = {
        "filename": os.path.basename(filepath),
        "sections": [],
        "headings": [],
        "tables": [],
        "images": [],
        "toc": {
            "title": None,
            "entries": [],
        },
        "formatting": {
            "default_font": None,
            "default_size": None,
            "page_margins": {},
        },
        "raw_text": "",
        "statistics": {},
    }

    # Extract page margins
    try:
        section = doc.sections[0]
        result["formatting"]["page_margins"] = {
            "top": _emu_to_inches(section.top_margin),
            "bottom": _emu_to_inches(section.bottom_margin),
            "left": _emu_to_inches(section.left_margin),
            "right": _emu_to_inches(section.right_margin),
        }
    except Exception:
        pass

    # Extract default font
    try:
        style = doc.styles["Normal"]
        if style.font.name:
            result["formatting"]["default_font"] = style.font.name
        if style.font.size:
            result["formatting"]["default_size"] = style.font.size.pt
    except Exception:
        pass

    # Parse paragraphs into sections
    current_section = None
    all_text_lines = []
    current_page = 1
    
    # Heuristic page estimation
    # Calculate lines per page from margins and font size
    margins = result["formatting"].get("page_margins", {})
    top_margin = margins.get("top") or 1.0
    bottom_margin = margins.get("bottom") or 1.0
    default_size_pt = result["formatting"].get("default_size") or 11
    
    # A4/Letter = ~11 inches height. Usable height = 11 - top - bottom
    usable_height_inches = 11.0 - top_margin - bottom_margin
    line_height_inches = (default_size_pt * 1.15) / 72  # 1.15 line spacing
    lines_per_page = max(int(usable_height_inches / line_height_inches), 30) if line_height_inches > 0 else 45
    cumulative_lines = 0
    pages_from_breaks = set()  # Track which pages we detected via hard breaks

    for para_idx, para in enumerate(doc.paragraphs):
        # Check for page breaks (hard breaks or rendered breaks)
        has_page_break = False
        for run in para.runs:
            if 'lastRenderedPageBreak' in run._element.xml or 'w:br w:type="page"' in run._element.xml:
                has_page_break = True
        
        if has_page_break:
            current_page += 1
            cumulative_lines = 0
            pages_from_breaks.add(current_page)
        else:
            # Estimate lines taken by this paragraph
            text = para.text.strip()
            if text:
                # Estimate character width ~80 chars per line for typical doc
                est_lines = max(1, len(text) // 80 + 1)
                cumulative_lines += est_lines
                
                # If cumulative lines exceed page capacity, bump page
                if cumulative_lines >= lines_per_page:
                    current_page += 1
                    cumulative_lines = 0

        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        is_toc_entry = _is_toc_paragraph(style_name, para)

        if not result["toc"]["title"] and _is_toc_title(text, style_name):
            result["toc"]["title"] = text

        # Detect heading levels
        heading_level = _get_heading_level(para, style_name, text)
        if is_toc_entry:
            heading_level = None

        # Extract paragraph formatting
        para_format = _extract_paragraph_format(para)

        # Extract run-level formatting (font, size, bold, italic)
        runs_info = _extract_runs_info(para)

        para_data = {
            "index": para_idx,
            "text": text,
            "style": style_name,
            "heading_level": heading_level,
            "format": para_format,
            "runs": runs_info,
            "has_image": _paragraph_has_image(para),
            "page": current_page,
        }

        if is_toc_entry:
            toc_entry = _extract_toc_entry(text, style_name, para_idx, current_page)
            if toc_entry:
                result["toc"]["entries"].append(toc_entry)

        if heading_level:
            heading_number, heading_title = _split_heading_number(text, heading_level)
            result["headings"].append({
                "text": text,
                "number": heading_number,
                "title": heading_title,
                "level": heading_level,
                "style": style_name,
                "index": para_idx,
                "page": current_page,
            })

        if heading_level and heading_level <= 3:
            # Start new section
            if current_section:
                result["sections"].append(current_section)
            current_section = {
                "heading": text,
                "level": heading_level,
                "paragraphs": [para_data],
                "start_index": para_idx,
                "page": current_page,
            }
        elif current_section:
            current_section["paragraphs"].append(para_data)
        else:
            # Content before first heading
            if not current_section:
                current_section = {
                    "heading": "(Document Header / Preamble)",
                    "level": 0,
                    "paragraphs": [para_data],
                    "start_index": para_idx,
                    "page": current_page,
                }
            else:
                current_section["paragraphs"].append(para_data)

        if text:
            all_text_lines.append(text)

    if current_section:
        result["sections"].append(current_section)

    result["raw_text"] = "\n".join(all_text_lines)

    # Extract headers and footers
    header_footer_text = []
    try:
        for i, section in enumerate(doc.sections):
            if hasattr(section, 'header') and section.header:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        header_footer_text.append(text)
            if hasattr(section, 'footer') and section.footer:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        header_footer_text.append(text)
    except Exception as e:
        print(f"Error extracting headers/footers: {e}")

    if header_footer_text:
        result["sections"].append({
            "heading": "(Document Headers & Footers)",
            "level": 0,
            "paragraphs": [{"index": 0, "text": t, "style": "", "heading_level": None, "format": {}, "runs": [], "has_image": False, "page": 1} for t in header_footer_text],
            "start_index": 0,
            "page": 1,
        })
        result["raw_text"] += "\n" + "\n".join(header_footer_text)

    # Parse tables
    for tbl_idx, table in enumerate(doc.tables):
        table_data = _extract_table(table, tbl_idx)
        result["tables"].append(table_data)

    # Extract images
    result["images"] = _extract_images(doc, filepath)

    # Compute statistics
    result["statistics"] = _compute_statistics(result)

    return result


def parse_excel(filepath):
    """
    Parse an Excel (.xlsx/.xls) file and extract content in a structured format
    fully compatible with the Word parser output. Each sheet becomes a section
    AND a table entry so all review steps work correctly.
    """
    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
        sections = []
        tables = []
        total_words = 0
        all_text_lines = []

        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            ws = wb[sheet_name]
            
            # ---- Build rows data (skip fully empty rows) ----
            sheet_rows = []
            num_cols = 0
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                if not any(c.strip() for c in cells):
                    continue
                num_cols = max(num_cols, len(cells))
                sheet_rows.append(cells)
            
            if not sheet_rows:
                continue  # skip blank sheets
            
            # ---- Build section (for chunk-based text review) ----
            paragraphs = []
            # Sheet heading
            paragraphs.append({
                "text": f"SHEET: {sheet_name}",
                "heading_level": 1,
                "alignment": "LEFT",
                "page": sheet_idx + 1,
                "runs": [{"text": f"SHEET: {sheet_name}", "bold": True, "italic": False, "size_pt": 14}],
                "has_image": False,
            })
            all_text_lines.append(f"=== SHEET: {sheet_name} ({len(sheet_rows)} rows × {num_cols} cols) ===")
            
            # Detect if the first row is a header row
            header_row = sheet_rows[0] if sheet_rows else []
            
            # Add each row as a paragraph with row number context
            for row_idx, row_cells in enumerate(sheet_rows):
                row_label = f"Row {row_idx + 1}"
                if row_idx == 0:
                    row_label = "Header Row"
                row_text = f"[{row_label}] {' | '.join(row_cells)}"
                word_count = len(row_text.split())
                total_words += word_count
                
                paragraphs.append({
                    "text": row_text,
                    "heading_level": None,
                    "alignment": "LEFT",
                    "page": sheet_idx + 1,
                    "runs": [{"text": row_text, "bold": (row_idx == 0), "italic": False, "size_pt": 10}],
                    "has_image": False,
                })
                all_text_lines.append(row_text)
            
            sections.append({
                "heading": f"Sheet: {sheet_name}",
                "level": 1,
                "page": sheet_idx + 1,
                "paragraphs": paragraphs,
            })
            
            # ---- Build table entry (for table-specific review) ----
            tables.append({
                "index": sheet_idx,
                "name": f"Sheet: {sheet_name}",
                "num_rows": len(sheet_rows),
                "num_cols": num_cols,
                "rows": [row_cells for row_cells in sheet_rows],
            })

        # ---- Build formatting stub (required by get_document_summary) ----
        formatting = {
            "default_font": None,
            "default_size": None,
            "page_margins": {},
        }

        parsed = {
            "filename": os.path.basename(filepath),
            "sections": sections,
            "tables": tables,
            "images": [],
            "formatting": formatting,
            "statistics": {
                "total_words": total_words,
                "total_pages": len(wb.sheetnames),
                "total_characters": sum(len(line) for line in all_text_lines),
                "total_sections": len(sections),
                "total_tables": len(tables),
                "total_images": 0,
                "empty_sections": 0,
            },
            "metadata": {"source_type": "excel"},
            "raw_text": "\n".join(all_text_lines),
        }

        return parsed

    except Exception as e:
        return {
            "filename": os.path.basename(filepath),
            "statistics": {
                "total_words": 0, "total_pages": 0, "total_characters": 0,
                "total_sections": 0, "total_tables": 0, "total_images": 0, "empty_sections": 0,
            },
            "formatting": {"default_font": None, "default_size": None, "page_margins": {}},
            "sections": [{"heading": "Error", "level": 1, "page": 1,
                          "paragraphs": [{"text": f"Error parsing Excel: {str(e)}", "heading_level": None,
                                          "alignment": "LEFT", "page": 1, "runs": [], "has_image": False}]}],
            "tables": [],
            "images": [],
            "raw_text": "",
        }



def _emu_to_inches(emu_val):
    """Convert EMU to inches."""
    if emu_val is None:
        return None
    return round(emu_val / 914400, 2)


def _get_heading_level(para, style_name, text):
    """Determine heading level from style or text pattern."""
    if not style_name:
        return None

    if _is_toc_style(style_name):
        return None

    # Built-in heading styles
    if style_name.startswith("Heading"):
        try:
            return int(style_name.replace("Heading ", "").replace("Heading", "").strip())
        except ValueError:
            return 1

    # Detect numbered section headings by pattern (e.g., "3.1.2 Title")
    if _looks_like_numbered_heading(para, text):
        # Count the depth by number of dots
        match = re.match(r"^(\d+(\.\d+)*)", text)
        if match:
            parts = match.group(1).split(".")
            return min(len(parts), 6)

    return None


def _looks_like_numbered_heading(para, text):
    """Heuristic fallback for numbered headings when proper heading styles are missing."""
    if not text:
        return False

    normalized = re.sub(r"\s+", " ", text).strip()
    if len(normalized) > 140:
        return False
    if normalized.endswith((".", ";", ":")):
        return False
    if not re.match(r"^\d+(?:\.\d+)*\.?\s+[A-Z]", normalized):
        return False

    try:
        ppr = para._element.pPr
        if ppr is not None and ppr.numPr is not None:
            return False
    except Exception:
        pass

    return True


def _is_toc_style(style_name):
    """Return True when the paragraph style indicates a TOC entry."""
    if not style_name:
        return False
    normalized = style_name.strip().lower()
    return bool(re.match(r"^toc\s*\d+$", normalized) or re.match(r"^toc\d+$", normalized))


def _is_toc_paragraph(style_name, para):
    """Detect TOC entry paragraphs produced by Word or styled manually."""
    if _is_toc_style(style_name):
        return True

    raw_text = para.text or ""
    if not _looks_like_toc_line(raw_text):
        return False

    try:
        xml = para._element.xml
    except Exception:
        return False

    return "_Toc" in xml and ("w:hyperlink" in xml or "PAGEREF" in xml)


def _is_toc_title(text, style_name):
    """Detect a TOC title line such as 'Table of Contents' or 'Contents'."""
    if not text:
        return False

    normalized = re.sub(r"\s+", " ", text).strip().lower()
    if normalized in {"table of contents", "contents"}:
        return True

    return bool(style_name and style_name.strip().lower() == "toc heading")


def _extract_toc_entry(text, style_name, para_idx, page):
    """Extract structured TOC entry data from a TOC paragraph."""
    clean_text = re.sub(r"\s+", " ", text.replace("\t", " ")).strip()
    if not clean_text:
        return None

    level = _get_toc_level(style_name)
    line_match = re.match(r"^(?P<label>.+?)(?:\.{2,}|\s{2,}|\t+)\s*(?P<page>\d+)\s*$", text.strip())
    if not line_match:
        line_match = re.match(r"^(?P<label>.+?)\s+(?P<page>\d+)\s*$", clean_text)

    label = clean_text
    page_ref = None
    if line_match:
        label = re.sub(r"\s+", " ", line_match.group("label")).strip()
        page_ref = line_match.group("page")

    if _is_toc_title(label, style_name):
        return None

    number, title = _split_heading_number(label, level)
    return {
        "text": label,
        "number": number,
        "title": title,
        "level": level,
        "page_ref": page_ref,
        "index": para_idx,
        "page": page,
    }


def _get_toc_level(style_name):
    """Extract the TOC nesting level from styles such as 'TOC 2'."""
    if not style_name:
        return 1

    match = re.search(r"(\d+)$", style_name.strip())
    if match:
        return int(match.group(1))
    return 1


def _looks_like_toc_line(text):
    """Return True when text has a typical TOC leader and page-number pattern."""
    if not text:
        return False

    stripped = text.strip()
    return bool(re.search(r"(?:\t+|\.{2,}|\s{2,})\s*\d+\s*$", stripped))


def _split_heading_number(text, level=None):
    """Split a heading into section number and title."""
    if not text:
        return None, ""

    normalized = re.sub(r"\s+", " ", text).strip()
    match = re.match(r"^(?P<number>\d+(?:\.\d+)*)\.?\s+(?P<title>.+)$", normalized)
    if match:
        number = match.group("number")
        if "." not in number and level and level > 1:
            return None, normalized
        return number, match.group("title").strip()

    return None, normalized


def _extract_paragraph_format(para):
    """Extract formatting details from a paragraph."""
    fmt = {}
    pf = para.paragraph_format

    # Alignment
    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    if pf.alignment is not None:
        fmt["alignment"] = alignment_map.get(pf.alignment, str(pf.alignment))

    # Spacing
    if pf.space_before is not None:
        fmt["space_before"] = str(pf.space_before)
    if pf.space_after is not None:
        fmt["space_after"] = str(pf.space_after)
    if pf.line_spacing is not None:
        fmt["line_spacing"] = str(pf.line_spacing)

    # Indentation
    if pf.first_line_indent is not None:
        fmt["first_line_indent"] = str(pf.first_line_indent)
    if pf.left_indent is not None:
        fmt["left_indent"] = str(pf.left_indent)

    return fmt


def _extract_runs_info(para):
    """Extract run-level formatting information."""
    runs = []
    for run in para.runs:
        run_info = {
            "text": run.text,
        }
        if run.font.name:
            run_info["font"] = run.font.name
        if run.font.size:
            run_info["size_pt"] = run.font.size.pt
        if run.bold:
            run_info["bold"] = True
        if run.italic:
            run_info["italic"] = True
        if run.underline:
            run_info["underline"] = True
        if run.font.color and run.font.color.rgb:
            run_info["color"] = str(run.font.color.rgb)
        runs.append(run_info)
    return runs


def _paragraph_has_image(para):
    """Check if a paragraph contains an embedded image."""
    for run in para.runs:
        if run._element.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        ):
            return True
        if run._element.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
        ):
            return True
    # Also check inline shapes
    xml = para._element.xml
    if "blipFill" in xml or "v:imagedata" in xml:
        return True
    return False


def _extract_table(table, tbl_idx):
    """Extract table content and check formatting."""
    name = f"Table {tbl_idx + 1}"
    try:
        prev_element = table._element.getprevious()
        if prev_element is not None and prev_element.tag.endswith('p'):
            p = Paragraph(prev_element, table._parent)
            text = p.text.strip()
            if text:
                name = text[:150] # Use preceding text as table name, capped to prevent massive names
    except Exception:
        pass

    rows_data = []
    for row_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            cells.append(cell.text.strip())
        rows_data.append(cells)

    # Detect if it has a header row
    has_header = len(rows_data) > 1

    return {
        "index": tbl_idx,
        "name": name,
        "rows": rows_data,
        "num_rows": len(rows_data),
        "num_cols": len(rows_data[0]) if rows_data else 0,
        "has_header": has_header,
    }


def _extract_images(doc, filepath):
    """Extract images from the document with metadata."""
    images = []
    try:
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_data))
                    width, height = img.size
                    img_format = img.format or "unknown"

                    # Create a capped full-size version for AI Vision (max 1024x1024)
                    ai_img = img.copy()
                    if ai_img.mode != 'RGB':
                        ai_img = ai_img.convert('RGB')
                    ai_img.thumbnail((1024, 1024))
                    ai_buffer = io.BytesIO()
                    ai_img.save(ai_buffer, format="JPEG")
                    full_b64 = base64.b64encode(ai_buffer.getvalue()).decode("utf-8")

                    # Create a small thumbnail for reference
                    thumb_size = (200, 200)
                    img.thumbnail(thumb_size)
                    thumb_buffer = io.BytesIO()
                    img.save(thumb_buffer, format="PNG")
                    thumb_b64 = base64.b64encode(thumb_buffer.getvalue()).decode("utf-8")

                    images.append({
                        "rel_id": rel_id,
                        "width": width,
                        "height": height,
                        "format": img_format,
                        "size_bytes": len(image_data),
                        "thumbnail_b64": thumb_b64,
                        "full_b64": full_b64,
                        "is_small": width < 300 or height < 300,
                        "is_very_large": width > 3000 or height > 3000,
                    })
                except Exception as e:
                    images.append({
                        "rel_id": rel_id,
                        "error": str(e),
                    })
    except Exception:
        pass
    return images


def _compute_statistics(parsed):
    """Compute document statistics."""
    raw = parsed["raw_text"]
    words = raw.split()
    sections = parsed["sections"]

    return {
        "total_paragraphs": sum(len(s["paragraphs"]) for s in sections),
        "total_words": len(words),
        "total_characters": len(raw),
        "total_sections": len(sections),
        "total_tables": len(parsed["tables"]),
        "total_images": len(parsed["images"]),
        "empty_sections": sum(
            1 for s in sections
            if all(not p["text"] for p in s["paragraphs"])
        ),
    }


def get_document_summary(parsed):
    """
    Build a concise text summary of the document for LLM review.
    Includes structure, content, tables, and formatting notes.
    """
    lines = []
    lines.append(f"# Document: {parsed['filename']}")
    lines.append(f"Total Words: {parsed['statistics']['total_words']}")
    lines.append(f"Total Sections: {parsed['statistics']['total_sections']}")
    lines.append(f"Total Tables: {parsed['statistics']['total_tables']}")
    lines.append(f"Total Images: {parsed['statistics']['total_images']}")

    # Formatting metadata
    fmt = parsed.get("formatting", {})
    if fmt.get("default_font"):
        lines.append(f"Default Font: {fmt.get('default_font')}")
    if fmt.get("default_size"):
        lines.append(f"Default Font Size: {fmt.get('default_size')}pt")
    margins = fmt.get("page_margins", {})
    if margins:
        lines.append(f"Page Margins (inches): Top={margins.get('top')}, Bottom={margins.get('bottom')}, Left={margins.get('left')}, Right={margins.get('right')}")

    lines.append("")

    # Sections with content
    for section in parsed["sections"]:
        heading = section["heading"]
        level = section["level"]
        page = section.get("page", 1)
        prefix = "#" * max(level, 1) if level else "##"
        lines.append(f"\n{prefix} {heading} (Page {page})")

        for para in section["paragraphs"]:
            text = para["text"]
            if not text:
                continue

            # Note formatting issues
            notes = []
            if para["has_image"]:
                notes.append("[CONTAINS IMAGE]")

            # Check font consistency
            for run in para.get("runs", []):
                if "font" in run and fmt["default_font"] and run["font"] != fmt["default_font"]:
                    notes.append(f"[FONT MISMATCH: {run['font']} vs default {fmt['default_font']}]")
                    break
                if "size_pt" in run and fmt["default_size"] and run["size_pt"] != fmt["default_size"]:
                    if para["heading_level"] is None:  # Only flag non-headings
                        notes.append(f"[SIZE MISMATCH: {run['size_pt']}pt vs default {fmt['default_size']}pt]")
                        break

            note_str = " ".join(notes) if notes else ""
            lines.append(f"  {text[:500]} {note_str}")

    # Tables
    if parsed["tables"]:
        lines.append("\n## TABLES IN DOCUMENT")
        for tbl in parsed["tables"]:
            lines.append(f"\n### Table {tbl['index'] + 1} ({tbl['num_rows']}×{tbl['num_cols']})")
            for row_idx, row in enumerate(tbl["rows"][:20]):  # First 20 rows instead of 5
                lines.append(f"  Row {row_idx}: {' | '.join(r[:60] for r in row)}")
            if tbl["num_rows"] > 20:
                lines.append(f"  ... ({tbl['num_rows'] - 20} more rows)")

    # Images
    if parsed["images"]:
        lines.append("\n## IMAGES IN DOCUMENT")
        for img in parsed["images"]:
            if "error" in img:
                lines.append(f"  Image {img['rel_id']}: ERROR - {img['error']}")
            else:
                size_note = ""
                if img.get("is_small"):
                    size_note = " [WARNING: SMALL IMAGE - may be unreadable]"
                if img.get("is_very_large"):
                    size_note = " [NOTE: Very large image]"
                lines.append(
                    f"  Image {img['rel_id']}: {img['width']}x{img['height']} "
                    f"{img['format']} ({img['size_bytes']//1024}KB){size_note}"
                )

    return "\n".join(lines)


def get_section_chunks(parsed, max_chars=6000):
    """
    Split the document into chunks suitable for LLM processing.
    Each chunk contains one or more sections, staying under max_chars.
    """
    chunks = []
    current_chunk = []
    current_size = 0

    for section in parsed["sections"]:
        section_text = _section_to_text(section)
        section_size = len(section_text)

        if current_size + section_size > max_chars and current_chunk:
            chunks.append("\n\n".join(current_chunk))
            current_chunk = []
            current_size = 0

        current_chunk.append(section_text)
        current_size += section_size

    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    return chunks


def _section_to_text(section):
    """Convert a section dict to readable text with paragraph index markers."""
    lines = [f"[Section: {section['heading']} (Starts on Page {section.get('page', 1)})]"]
    current_page = None
    for para in section["paragraphs"]:
        if para["text"]:
            para_page = para.get("page", 1)
            para_idx = para.get("index", 0)
            if current_page != para_page:
                lines.append(f"\n--- [Page {para_page}] ---")
                current_page = para_page
            # Include paragraph index marker for precise error location
            lines.append(f"[¶{para_idx}] {para['text']}")
    return "\n".join(lines)


def _sanitize_docx(filepath):
    """
    Sanitize a corrupt docx file by removing relationships pointing to 'NULL'.
    This addresses the common 'word/NULL' archive KeyErrors on python-docx
    caused by third-party document generators handling images improperly.
    """
    tmp_path = filepath + ".tmp"
    with zipfile.ZipFile(filepath, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w') as zout:
            for item in zin.infolist():
                try:
                    content = zin.read(item.filename)
                except KeyError:
                    continue  # Skip unreadable items
                
                # We only care about Relationship XML files
                if item.filename.endswith('.rels'):
                    content_str = content.decode('utf-8', errors='ignore')
                    if 'Target="NULL"' in content_str:
                        # Remove any Relationship nodes that target 'NULL'
                        content_str = re.sub(r'<Relationship[^>]*?Target="NULL"[^>]*?/>', '', content_str)
                        content = content_str.encode('utf-8')
                zout.writestr(item, content)
                
    # Replace the original file with the sanitized version
    shutil.move(tmp_path, filepath)
