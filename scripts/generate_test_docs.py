from docx import Document
import os

def create_error_doc(path):
    doc = Document()
    doc.add_heading('System Design Specification v2.0', 0)
    
    doc.add_heading('1. Introduction', 1)
    doc.add_paragraph("This document outlines the arcitecture of the new billing system. The total project timeline is set for exactly 6 months starting from Q1.")
    
    doc.add_heading('2. Database Design', 1)
    doc.add_paragraph("We have decided to use MongoDB as our primary relational database because it enforces strict SQL schemas perfectly.")
    
    doc.add_heading('3. API Endpoints', 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Endpoint URL'
    hdr_cells[1].text = 'HTTP Method'
    hdr_cells[2].text = 'Description'
    
    row_cells = table.add_row().cells
    row_cells[0].text = '/api/v2/login'
    row_cells[1].text = 'PST'
    row_cells[2].text = 'Authenticates the user using JWT.'
    
    row_cells = table.add_row().cells
    row_cells[0].text = '/api/v2/invoices'
    row_cells[1].text = 'GETT'
    row_cells[2].text = 'Fetches the list of all invoices.'
    
    doc.add_heading('4. Conclusion', 1)
    doc.add_paragraph("The system will be fully deployed and completed in 3 months, as stated in the introduction.")
    
    doc.save(path)

def create_fixed_doc(path):
    doc = Document()
    doc.add_heading('System Design Specification v2.0', 0)
    
    doc.add_heading('1. Introduction', 1)
    doc.add_paragraph("This document outlines the architecture of the new billing system. The total project timeline is set for exactly 6 months starting from Q1.")
    
    doc.add_heading('2. Database Design', 1)
    doc.add_paragraph("We have decided to use PostgreSQL as our primary relational database because it enforces strict SQL schemas perfectly.")
    
    doc.add_heading('3. API Endpoints', 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Endpoint URL'
    hdr_cells[1].text = 'HTTP Method'
    hdr_cells[2].text = 'Description'
    
    row_cells = table.add_row().cells
    row_cells[0].text = '/api/v2/login'
    row_cells[1].text = 'POST'
    row_cells[2].text = 'Authenticates the user using JWT.'
    
    row_cells = table.add_row().cells
    row_cells[0].text = '/api/v2/invoices'
    row_cells[1].text = 'GET'
    row_cells[2].text = 'Fetches the list of all invoices.'
    
    doc.add_heading('4. Conclusion', 1)
    doc.add_paragraph("The system will be fully deployed and completed in 6 months, as stated in the introduction.")
    
    doc.save(path)

if __name__ == "__main__":
    error_path = r'C:\Users\yash badgujar\Downloads\TICO\Test_Document_With_Errors.docx'
    fixed_path = r'C:\Users\yash badgujar\Downloads\TICO\Test_Document_Fixed.docx'
    
    create_error_doc(error_path)
    create_fixed_doc(fixed_path)
    print(f"Generated documents at:\n1. {error_path}\n2. {fixed_path}")
